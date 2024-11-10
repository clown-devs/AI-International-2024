import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from parser.parser import Analytics

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

def save_analytics_to_word(analytics: Analytics, output_file: str) -> None:
    print("[DEBUG] Saving analytics to word")
    doc = Document()
    
    # Заголовок отчета
    doc.add_heading('Отчет по анализу кортикограмм', level=1)
    
    # Общая аналитическая информация
    doc.add_paragraph(f'Общее количество аномалий: {analytics.anomaly_count}')
    doc.add_paragraph(f'Общая продолжительность записи (сек): {analytics.total_time:.2f}')
    doc.add_paragraph(f'Общая длительность аномалий (сек): {analytics.total_duration:.2f}')
    doc.add_paragraph(f'Средняя длительность аномалии (сек): {analytics.average_duration:.2f}')
    doc.add_paragraph(f'Процент времени с аномалиями: {analytics.time_with_anomalies:.2f}%')
    doc.add_paragraph(f'Средний интервал между аномалиями (сек): {analytics.average_interval:.2f}')
    
    # Заголовок таблицы
    # Аномалии в таблице указаны только если их длительность больше 5 секунд
    #doc.add_paragraph('В таблице указаны только аномалии, длительность которых больше 5 секунд')
    doc.add_heading('Детали аномалий', level=2)
    
    # Создаем таблицу
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Заголовок таблицы
    hdr_cells = table.rows[0].cells
    headers = ['Тип аномалии', 'Начало (сек)', 'Конец (сек)', 'Пиковая амплитуда']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[i]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="BFBFBF"/>'.format(nsdecls('w'))))

    # Объединяем и сортируем все аномалии по времени начала
    sorted_anomalies = []
    for anomaly_type, intervals in analytics.anomalies_by_type.items():
        for i, (start, end) in enumerate(intervals):
            anomaly_name = 'SWD' if anomaly_type == 1 else 'IS' if anomaly_type == 2 else 'DS'
            peak_amplitude = analytics.peak_amplitudes[anomaly_type][i]
            sorted_anomalies.append((anomaly_name, start, end, peak_amplitude))

    sorted_anomalies.sort(key=lambda x: x[1]) 
    
    for i, (anomaly_name, start, end, peak_amplitude) in enumerate(sorted_anomalies):
        # если меньше 5 секунд то не добавляем
        #if end - start < 5:
        #    continue
        row_cells = table.add_row().cells
        row_cells[0].text = anomaly_name
        row_cells[1].text = f'{start:.2f}'
        row_cells[2].text = f'{end:.2f}'
        row_cells[3].text = f'{peak_amplitude:.10e}'

        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

        shading_color = "D3D3D3" if i % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shading_color)))
    doc.save(output_file)
    print ("[DEBUG] Word file saved to", output_file)
